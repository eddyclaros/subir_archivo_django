from re import T
from django.shortcuts import render
from django.http import HttpResponse
import io
from docx import Document
from docx.shared import Inches,Pt,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK

from generarword.models import Clausulas

def inicio(request):
    return HttpResponse("<h1>Bienvenido</h1>")

def principal(request):
    return render(request,'paginas/principal.html')

def publication(request):
    return render(request,'paginas/publication.html')

def coverletter_export(request):
    clausula=Clausulas.objects.all()
    print(clausula)
    titulo=clausula
    document = Document()
    for clau in clausula:
        print(clau.title)
        paragraph=document.add_paragraph(clau.title).bold=True
        paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

        paragraph=document.add_paragraph(clau.subtitle).bold=True
        paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

        paragraph=document.add_paragraph(clau.numpoliza +" AUTI-00001")
        paragraph=document.add_paragraph(clau.nameaseg +" JUAN PEREZ")

        paragraph=document.add_paragraph(clau.textpoliza)
        paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
        


    #document=Document()
    
    # Añadimos un titulo al documento, a nivel 0
    #document.add_heading('Documento creado con Python', 1)
    #document.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    # Añadimos un párrafo
    
    #document.add_paragraph()


    """ paragraph=document.add_paragraph("Queda entendido y convenido que, no obstante, lo que se establece en las Condiciones Generales y Particulares de la póliza que, la Compañía hará efectivo un pago parcial al Asegurado equivalente al 50% (Cincuenta por ciento) del siniestro, siempre y cuando se hayan dado cumplimiento a las siguientes condiciones:")
    paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    p = document.add_paragraph('El contenido de los párrafos se añadir en varias líneas. ')
    p.add_run('Pudiéndose configurar que el texto tenga formato tipo ')
    p.add_run('negrita').bold = True
    p.add_run(' o ')
    p.add_run('itálica.').italic = True
    # Para indicar subtitulo se indica el nivel 1
    document.add_heading('Subtitulo', level=1)
    document.add_paragraph('Ahora se puede crear una enumeración')
    document.add_paragraph('Uno', style='List Number')
    document.add_paragraph('Dos', style='List Number')
    document.add_paragraph('Tres', style='List Number')
    document.add_paragraph('O viñetas')
    document.add_paragraph('Manzana', style='List Bullet') 
    document.add_paragraph('Pera', style='List Bullet')
    document.add_paragraph('Naranja', style='List Bullet')
    # Imágenes
    #document.add_heading('Imágenes', level=1)
    #document.add_picture('analytics_lane.jpg', width=Cm(5))
    # Tablas
    document.add_heading('Tablas', level=1)
    data = (('Manzana', 12), ('Pera', 5), ('Naranja', 12))
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'Fruta'
    table.rows[0].cells[1].text = 'Cantidad'
    for prod, numbr in data:
        row_cells = table.add_row().cells
        row_cells[0].text = prod
        row_cells[1].text = str(numbr) """
    document.save('ejemplo.docx')

    document_data = io.BytesIO()
    document.save(document_data)
    document_data.seek(0)
    response = HttpResponse(
        document_data.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = 'attachment; filename = "Certificado.docx"'
    response["Content-Encoding"] = "UTF-8"
    return response

    


